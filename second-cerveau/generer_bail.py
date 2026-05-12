#!/usr/bin/env python3
"""
Génération d'un bail meublé (DOCX + PDF) à partir d'une fiche locataire du vault
et du référentiel `bail-config.yml`.

Usage:
    python generer_bail.py --locataire "Hella Taoutaou" --date-debut 2024-05-23
    python generer_bail.py --locataire "Milo Rouille" --date-debut 2026-04-11 \
        --date-signature 2026-04-10 --loyer 600 --charges 90
    python generer_bail.py --locataire "Kenan" --date-debut 2023-08-23 --dry-run

Sortie:
    G:\\Mon Drive\\Obsidian\\08. Outils\\Baux\\_generes\\<Locataire>\\
        Bail-<Bien>-<Locataire>-<YYYY-MM-DD>.docx
        Bail-<Bien>-<Locataire>-<YYYY-MM-DD>.pdf   (si conversion possible)

Pré-requis:
    pip install -r requirements.txt
"""

from __future__ import annotations
import argparse
import datetime as dt
import os
import re
import sys
import unicodedata
from dataclasses import dataclass
from pathlib import Path

import yaml
from num2words import num2words
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Configuration (chemins)
# ---------------------------------------------------------------------------

VAULT_ROOT = Path(os.environ.get("VAULT_ROOT", r"G:\Mon Drive\Obsidian"))
LOCATAIRES_DIR_ACTUELS = VAULT_ROOT / "07. Contacts" / "05. Locataires" / "01. Actuels"
LOCATAIRES_DIR_CANDIDATS = VAULT_ROOT / "07. Contacts" / "05. Locataires" / "_Candidats"
BAUX_DIR = VAULT_ROOT / "08. Outils" / "Baux"
QUITTANCES_DIR = VAULT_ROOT / "08. Outils" / "Quittances"
BAIL_CONFIG = BAUX_DIR / "bail-config.yml"
BIENS_CONFIG = QUITTANCES_DIR / "biens.yml"
OUTPUT_DIR = Path(os.environ.get("OUTPUT_DIR_BAUX", str(BAUX_DIR / "_generes")))

MOIS_FR = [
    "", "janvier", "février", "mars", "avril", "mai", "juin",
    "juillet", "août", "septembre", "octobre", "novembre", "décembre"
]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def date_en_lettres(d: dt.date) -> str:
    """Convertit 2026-04-11 → 'onze avril deux mille vingt-six'."""
    jour = num2words(d.day, lang="fr")
    mois = MOIS_FR[d.month]
    annee = num2words(d.year, lang="fr")
    return f"{jour} {mois} {annee}"


def date_chiffres(d: dt.date) -> str:
    """2026-04-11 → '11 avril 2026' (avec '1er' pour le 1er)."""
    j = "1er" if d.day == 1 else str(d.day)
    return f"{j} {MOIS_FR[d.month]} {d.year}"


def montant_en_lettres(montant: float) -> str:
    n = int(round(montant))
    return num2words(n, lang="fr")


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class Locataire:
    nom_fichier: str
    nom_affiche: str
    civilite: str | None         # "Monsieur" / "Madame" / "Mademoiselle"
    date_naissance: dt.date | None
    lieu_naissance: str | None
    nationalite: str | None
    email: str | None
    adresse_bien: str
    surface_m2: int | None       # surface réelle du logement (prioritaire sur bail-config)
    montant_loyer: float
    montant_charges: float
    depot_garantie: float | None
    date_entree_bail: dt.date | None
    jour_paiement: int | None    # 1, 2, 5...

    @property
    def civilite_abregee(self) -> str:
        c = (self.civilite or "").lower()
        if c.startswith("mons"):
            return "M."
        if c.startswith("mada"):
            return "Mme"
        if c.startswith("made") or c == "melle":
            return "Mlle"
        return "M./Mme"   # fallback explicite

    @property
    def est_feminin(self) -> bool:
        c = (self.civilite or "").lower()
        return c.startswith("mada") or c.startswith("made") or c == "melle"


# ---------------------------------------------------------------------------
# Lecture des données
# ---------------------------------------------------------------------------

FRONTMATTER_RE = re.compile(r"^---\n(.*?)\n---\n", re.DOTALL)

def _parse_date(v):
    if v is None or v == "":
        return None
    if isinstance(v, dt.date):
        return v
    try:
        return dt.date.fromisoformat(str(v))
    except ValueError:
        return None


def lire_fiche_locataire(path: Path) -> Locataire | None:
    text = path.read_text(encoding="utf-8")
    m = FRONTMATTER_RE.match(text)
    if not m:
        return None
    fm = yaml.safe_load(m.group(1)) or {}
    return Locataire(
        nom_fichier=path.stem,
        nom_affiche=fm.get("nom_officiel") or path.stem,
        civilite=fm.get("civilite") or None,
        date_naissance=_parse_date(fm.get("date_naissance")),
        lieu_naissance=fm.get("lieu_naissance") or None,
        nationalite=fm.get("nationalite") or None,
        email=fm.get("email") or None,
        adresse_bien=fm.get("adresse_bien") or "",
        surface_m2=int(fm["surface_m2"]) if fm.get("surface_m2") else None,
        montant_loyer=float(fm.get("montant_loyer") or 0),
        montant_charges=float(fm.get("montant_charges") or 0),
        depot_garantie=float(fm["depot_garantie"]) if fm.get("depot_garantie") else None,
        date_entree_bail=_parse_date(fm.get("date_entree_bail")),
        jour_paiement=int(fm["jour_paiement"]) if fm.get("jour_paiement") else None,
    )


def trouver_locataire(query: str) -> tuple[Locataire, str]:
    """Cherche dans 01. Actuels/ puis _Candidats/. Retourne (locataire, statut).
    statut = 'actuel' ou 'candidat'."""
    def norm(s: str) -> str:
        s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode().lower()
        return re.sub(r"\s+", " ", s).strip()
    qn = norm(query)
    matches = []
    for folder, statut in [(LOCATAIRES_DIR_ACTUELS, "actuel"), (LOCATAIRES_DIR_CANDIDATS, "candidat")]:
        if not folder.exists():
            continue
        for f in folder.glob("*.md"):
            if f.name.startswith("_"):
                continue
            if qn in norm(f.stem):
                loc = lire_fiche_locataire(f)
                if loc:
                    matches.append((loc, statut))
    if not matches:
        raise SystemExit(f"Locataire introuvable (ni dans Actuels ni dans _Candidats) : {query}")
    if len(matches) > 1:
        noms = ", ".join(f"{m[0].nom_fichier} ({m[1]})" for m in matches)
        raise SystemExit(f"Plusieurs matchs : {noms}. Préciser.")
    return matches[0]


def charger_config():
    with open(BAIL_CONFIG, "r", encoding="utf-8") as f:
        bail_cfg = yaml.safe_load(f)
    with open(BIENS_CONFIG, "r", encoding="utf-8") as f:
        biens_cfg = yaml.safe_load(f)
    return bail_cfg, biens_cfg


def resoudre_bien(loc: Locataire, biens_cfg, bail_cfg) -> dict:
    """Croise adresse_bien locataire + biens.yml + bail-config.yml.
    Retourne un dict consolidé avec adresse + complément + caractéristiques."""
    adresse_lc = loc.adresse_bien.lower()
    for bien in biens_cfg["biens"]:
        for frag in bien["match_adresse"]:
            if frag.lower() in adresse_lc:
                bid = bien["id"]
                bail_bien = bail_cfg["biens"].get(bid, {})
                # Complément (RDC, étage, etc.)
                if "description_complement" in bail_bien:
                    complement = bail_bien["description_complement"]
                elif "description_complement_template" in bail_bien:
                    m = re.search(r"(?:studio|appartement|apt)\s+(\d+)", loc.adresse_bien, re.IGNORECASE)
                    num = m.group(1) if m else "?"
                    complement = bail_bien["description_complement_template"].format(studio_num=num)
                elif "description_complement_par_cote" in bail_bien:
                    complement = bail_bien.get("description_complement_defaut", "RDC")
                    for cle, val in bail_bien["description_complement_par_cote"].items():
                        if cle in adresse_lc:
                            complement = val
                            break
                else:
                    complement = ""
                return {
                    "id": bid,
                    "adresse_ligne1": bien["ligne1"],
                    "complement": complement,
                    "cp_ville": bien["cp_ville"],
                    "surface_m2": bail_bien.get("surface_m2", "?"),
                    "pieces": bail_bien.get("pieces", "[à compléter]"),
                    "charges_incluses": bail_bien.get("charges_incluses", "[à compléter]"),
                    "inventaire_type": bail_bien.get("inventaire_type"),
                }
    raise SystemExit(f"Bien introuvable pour {loc.nom_fichier} (adresse='{loc.adresse_bien}').")


# ---------------------------------------------------------------------------
# Génération DOCX
# ---------------------------------------------------------------------------

def _para(doc, text, bold=False, align=None, italic=False, size=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    return p


def _heading(doc, text):
    return _para(doc, text, bold=True)


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def _check_fiche(loc: Locataire) -> list[str]:
    """Renvoie la liste des champs manquants critiques."""
    issues = []
    if not loc.civilite:
        issues.append("civilite (Monsieur/Madame/Mademoiselle)")
    if not loc.date_naissance:
        issues.append("date_naissance (YYYY-MM-DD)")
    if not loc.lieu_naissance:
        issues.append("lieu_naissance")
    if not loc.nationalite:
        issues.append("nationalite")
    if loc.montant_loyer <= 0:
        issues.append("montant_loyer")
    if not loc.adresse_bien:
        issues.append("adresse_bien")
    return issues


def render_bail(loc: Locataire, bien: dict, bail_cfg, args) -> Document:
    cfg_bailleur = bail_cfg["bailleur"]
    cfg_def = bail_cfg["defaults"]

    # Résolution des valeurs (arg CLI > frontmatter > defaults)
    loyer = args.loyer if args.loyer is not None else loc.montant_loyer
    charges = args.charges if args.charges is not None else loc.montant_charges
    depot = args.depot if args.depot is not None else (loc.depot_garantie or cfg_def["depot_garantie"])
    jour_paiement = args.jour_paiement if args.jour_paiement else (loc.jour_paiement or cfg_def["jour_paiement_loyer"])
    delai_restitution = args.delai_restitution or cfg_def["delai_restitution_depot"]
    date_debut = dt.date.fromisoformat(args.date_debut)
    date_signature = dt.date.fromisoformat(args.date_signature) if args.date_signature else date_debut - dt.timedelta(days=1)

    # Bailleur — date naissance
    d_naiss_bailleur = dt.date.fromisoformat(cfg_bailleur["date_naissance"])

    # Texte adapté féminin/masculin
    accord_ne = "Née" if loc.est_feminin else "Né"
    accord_designe = "désignée" if loc.est_feminin else "désigné"

    # Jour de paiement avec 1er
    jp = "1er" if jour_paiement == 1 else str(jour_paiement)

    doc = Document()
    # Style par défaut Arial 11
    sty = doc.styles["Normal"]
    sty.font.name = "Arial"
    sty.font.size = Pt(11)
    # Marges
    for s in doc.sections:
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)

    # ===== Titre =====
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(cfg_def["type_bail"])
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph()

    _para(doc, "Entre les soussignés,")
    doc.add_paragraph()

    # ===== Bailleur =====
    _para(doc, f"M. {cfg_bailleur['nom_avec_capitales']},", bold=True)
    _para(
        doc,
        f"Né le {d_naiss_bailleur.strftime('%d/%m/%Y')} à {cfg_bailleur['lieu_naissance']}, "
        f"de nationalité {cfg_bailleur['nationalite']}, demeurant au "
        f"{cfg_bailleur['adresse']}, {cfg_bailleur['cp_ville']}"
    )
    _para(doc, "désigné ci-après « le bailleur »", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    _para(doc, "Et")
    doc.add_paragraph()

    # ===== Preneur =====
    _para(doc, f"{loc.civilite_abregee} {loc.nom_affiche},", bold=True)
    _para(
        doc,
        f"{accord_ne} le {loc.date_naissance.strftime('%d/%m/%Y')} à {loc.lieu_naissance}, "
        f"de nationalité {loc.nationalite}"
    )
    _para(doc, f"{accord_designe} ci-après « le preneur »", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    _para(doc, "Il a été convenu et arrêté ce qui suit :")
    doc.add_paragraph()
    _para(
        doc,
        "Le bailleur loue les locaux et équipements ci-après désignés au preneur "
        "qui les accepte aux conditions suivantes :"
    )
    doc.add_paragraph()

    # ===== Adresse du logement =====
    _heading(doc, "Adresse du logement donné en location")
    _para(doc, bien["adresse_ligne1"])
    if bien["complement"]:
        _para(doc, bien["complement"])
    _para(doc, bien["cp_ville"])
    doc.add_paragraph()
    # Surface : priorité à la fiche locataire, puis bail-config
    surface = loc.surface_m2 if loc.surface_m2 else bien.get("surface_m2", "?")
    _para(doc, f"Un studio de {surface}m² avec {bien['pieces']}, meublé.")
    doc.add_paragraph()

    # ===== État des lieux =====
    _heading(doc, "État des lieux contradictoire")
    _para(doc, "Un état des lieux contradictoire sera établi lors de la remise des clés au locataire et lors de la restitution de celles-ci.")
    _para(doc, "L'état des lieux sera obligatoirement annexé au présent contrat.")
    doc.add_paragraph()

    # ===== Inventaire intro =====
    _heading(doc, "Inventaire contradictoire")
    _para(doc, "La présente location étant consentie et acceptée en meublé, un inventaire contradictoire des meubles sera établi lors de la remise des clés au locataire et lors de la restitution de celles-ci. L'inventaire sera annexé au présent contrat. Le preneur sera responsable de toute détérioration ou perte pouvant survenir à ce mobilier.")
    doc.add_paragraph()

    # ===== Durée =====
    _heading(doc, "Durée")
    _para(
        doc,
        f"Le logement constitue la résidence principale du locataire. La présente "
        f"location est consentie et acceptée pour une durée de {cfg_def['duree_bail']} "
        f"qui commence à courir le {date_en_lettres(date_debut)} – {date_chiffres(date_debut)} "
        f"renouvelable ensuite par tacite reconduction et par période d'un an faute de congé préalable."
    )
    doc.add_paragraph()

    # ===== Forme du congé =====
    _heading(doc, "Forme du congé, durée du préavis")
    _para(doc, "Le présent contrat pourra être résilié :")
    _bullet(doc, f"par le preneur à tout moment, moyennant un délai de préavis de {cfg_def['preavis_locataire']} (partant de la date de réception de l'acte)")
    _bullet(doc, f"par le bailleur, à l'expiration du bail ou de chacun de ses renouvellements, moyennant un délai de préavis de {cfg_def['preavis_bailleur']}. Lorsque le logement constitue la résidence principale du locataire, le refus du renouvellement du bail par le bailleur doit être motivé soit par sa décision de reprendre ou de vendre le logement, soit par un motif légitime et sérieux, notamment l'inexécution par le locataire de l'une des obligations lui incombant, et notifié au preneur 3 mois avant l'expiration du bail.")
    _para(doc, "Le congé de location devra être signifié de part et d'autre par lettre recommandée avec accusé de réception ou par acte d'huissier.")
    _para(doc, "En cas de résiliation du preneur ou à l'expiration du bail, la notification de résiliation ou de fin de bail vaudra engagement formel de partir et renonciation à tout maintien dans les lieux, sans qu'il soit besoin de recourir à aucune formalité. Faute de libérer les lieux à la date convenue, la clause pénale incluse dans le présent contrat sera immédiatement applicable.")
    doc.add_paragraph()

    # ===== Loyer =====
    _heading(doc, "Montant et paiement du loyer")
    _para(
        doc,
        f"La présente location est consentie et acceptée moyennant un paiement mensuel "
        f"et d'avance de {montant_en_lettres(loyer)} euros ({int(loyer)} euros). "
        f"Il sera payable le {jp} de chaque mois."
    )
    doc.add_paragraph()

    # ===== Charges =====
    _heading(doc, "Charges")
    _para(
        doc,
        f"Les charges sont fixées forfaitairement à {montant_en_lettres(charges)} euros "
        f"({int(charges)} euros) par mois. Elles seront acquittées en même temps que le "
        f"loyer, c'est-à-dire mensuellement et d'avance et révisées chaque année aux "
        f"mêmes conditions que le loyer principal. Elles contiennent notamment "
        f"{bien['charges_incluses']}."
    )
    doc.add_paragraph()

    # ===== Indexation =====
    _heading(doc, "Indexation")
    _para(
        doc,
        "Le loyer pourra être indexé annuellement à la date anniversaire du contrat, "
        "sur la base de l'Indice de Référence des Loyers (IRL) publié chaque trimestre "
        "par l'INSEE (article 17-1 de la loi n° 89-462 du 6 juillet 1989). L'indice de "
        "référence est celui du trimestre précédant la signature du présent contrat. "
        "À chaque révision, le nouveau loyer sera calculé selon la formule : "
        "nouveau loyer = loyer en cours × (nouvel IRL / IRL de référence)."
    )
    doc.add_paragraph()

    # ===== Dépôt =====
    _heading(doc, "Dépôt de garantie")
    _bullet(doc, f"À titre de garantie de l'entière exécution de ses obligations le locataire verse, ce jour, un dépôt de garantie correspondant à la somme de {montant_en_lettres(depot)} euros ({int(depot)} euros).")
    _bullet(doc, f"Ce dépôt qui ne pourra excéder deux mois de loyer principal ne dispensera en aucun cas le locataire du paiement du loyer et des charges aux dates fixées. Il sera restitué dans un délai maximal de {delai_restitution} à compter du départ du locataire, déduction faite, le cas échéant, des sommes restant dues au bailleur et des paiements dont ce dernier pourrait être tenu pour responsable aux lieu et place du locataire. Le départ s'entend après complet déménagement et établissement de l'état des lieux et de l'inventaire contradictoire de sortie, résiliation des abonnements en cours tels qu'eau, électricité, gaz, téléphone, exécution des réparations locatives, paiement des taxes et impôts et remise des clés. À défaut de restitution du montant de garantie dans le délai prévu, le solde du dépôt de garantie restant dû au locataire après arrêté des comptes produira intérêt au taux légal au profit du locataire.")
    _bullet(doc, "Le dépôt de garantie ne sera pas révisable au cours de la présente location.")
    doc.add_paragraph()

    # ===== Obligations preneur =====
    _heading(doc, "Obligations du preneur")
    _para(doc, "Le preneur est tenu aux obligations suivantes :")
    for o in [
        "de payer le loyer et les charges récupérables aux termes convenus ;",
        "d'user paisiblement du logement suivant la destination qui lui a été donnée par le contrat de location (exclusivement habitation) ;",
        "de répondre des dégradations et pertes qui surviennent pendant la durée du contrat dans les locaux dont il a la jouissance exclusive, à moins qu'il ne prouve qu'elles ont eu lieu par un cas de force majeure, par faute du bailleur ou par le fait d'un tiers qu'il n'a pas introduit dans le logement ;",
        "de prendre à sa charge l'entretien courant du logement ainsi que l'ensemble des réparations locatives définies par décret au Conseil d'État, sauf si elles sont occasionnées par vétusté, malfaçon, vice de construction, cas fortuit ou de force majeure ;",
        "de ne pas transformer sans l'accord exprès et écrit du bailleur les locaux loués et les équipements ; le bailleur peut, si le locataire a méconnu cette obligation, exiger la remise en état des locaux ou conserver les transformations effectuées sans que le locataire puisse réclamer une indemnité pour les frais engagés ; le bailleur a toutefois la faculté d'exiger, aux frais du locataire, la remise immédiate des lieux en état lorsque les transformations mettent en péril le bon fonctionnement des équipements ou la sécurité du local ;",
        "de s'assurer contre les risques dont il doit répondre en sa qualité de locataire et d'en justifier lors de la remise des clés et ensuite chaque année ;",
        "de souffrir la réalisation par le bailleur des réparations urgentes et qui ne peuvent être différées jusqu'à la fin du contrat de location sans préjudice de l'application des dispositions de l'article 1724 du Code civil ;",
        "de laisser visiter, aussitôt le congé donné ou reçu, ou en cas de mise en vente, les locaux loués, deux heures par jour, les jours ouvrables ;",
        "de ne céder le contrat de location, ni de sous-louer, sauf accord exprès et écrit du bailleur ;",
        "de ne pas utiliser l'accès Internet à des fins de reproduction, de représentation, de mise à disposition ou de communication au public d'œuvres ou d'objets protégés par un droit d'auteur ou par un droit voisin, tels que des textes, images, photographies, œuvres musicales, œuvres audiovisuelles, logiciels et jeux vidéo, sans autorisation.",
    ]:
        _bullet(doc, o)
    doc.add_paragraph()

    # ===== Obligations bailleur =====
    _heading(doc, "Obligations du bailleur")
    _para(doc, "Le bailleur est tenu aux principales obligations suivantes :")
    for o in [
        "de délivrer au locataire le logement en bon état de réparations de toutes espèces et les équipements mentionnés au contrat de location en bon état de fonctionnement ;",
        "d'assurer la jouissance paisible du logement et de garantir le locataire contre les vices ou défauts qui en empêchent l'usage, quand même il ne les aurait pas connus lors de la conclusion du contrat de location, sans préjudice de l'application du second alinéa de l'article 1721 du Code civil ;",
        "d'entretenir les locaux en état de servir à l'usage prévu dans le contrat et d'y faire toutes les réparations nécessaires autres que locatives ;",
        "de remettre, lorsque le locataire en fait la demande, une quittance gratuitement.",
    ]:
        _bullet(doc, o)
    doc.add_paragraph()

    # ===== Clause résolutoire =====
    _heading(doc, "Clause résolutoire")
    for c in [
        "À défaut de paiement de tout ou partie du loyer ou des charges et un mois après commandement demeuré infructueux, le présent contrat sera résilié immédiatement et de plein droit et le bailleur pourra, dans le cas où le locataire ne quitterait pas les lieux, l'y contraindre par simple ordonnance de référé.",
        "Il est expressément convenu qu'en cas de paiement par chèque le loyer et les charges ne seront considérés comme réglés qu'après encaissement du chèque, la clause résolutoire pouvant être appliquée par le bailleur dans le cas où le chèque serait sans provision.",
        "Toute offre de paiement ou d'exécution après l'expiration du délai ci-dessus mentionné sera réputée nulle et non avenue et ne pourra faire obstacle à la résiliation de la présente location.",
        "À défaut de production par le locataire d'attestation couvrant ses risques locatifs et un mois après commandement resté infructueux, il sera fait application de la présente clause résolutoire.",
    ]:
        _bullet(doc, c)
    doc.add_paragraph()

    # ===== Clause pénale =====
    _heading(doc, "Clause pénale")
    for c in [
        "À titre de clause pénale, le preneur accepte entièrement et définitivement d'avoir à payer au bailleur une somme égale à 10% des sommes dues, sans que ce paiement puisse le dispenser du règlement des sommes impayées et du règlement intégral des frais nécessaires au recouvrement de ces sommes.",
        "Ladite clause pénale sera applicable dans un délai de quinze jours après mise en demeure de payer, le tout sans qu'il soit dérogé à la précédente clause résolutoire.",
        "En cas de retard dans la libération des lieux après réception du congé ou expiration du contrat, le preneur, quels que soient ses motifs, devra une astreinte par jour de retard calculée sur la base de trois fois le loyer journalier en cours à la date du départ. Cette indemnité n'ouvrira aucun droit de maintien dans les lieux au preneur, et elle sera acquise au bailleur à titre d'indemnité, à forfait, sans préjudice de tous dommages et intérêts.",
    ]:
        _bullet(doc, c)
    doc.add_paragraph()

    # ===== Élection de domicile =====
    _heading(doc, "Élection de domicile")
    _para(doc, "Pour l'exécution des présentes et de leur suite, le bailleur fait élection de domicile en sa demeure et le preneur dans les lieux loués.")
    doc.add_paragraph()

    # ===== Annexes obligatoires =====
    _heading(doc, "Annexes obligatoires")
    _para(doc, "Sont annexés au présent contrat conformément à la loi du 6 juillet 1989 et à ses textes d'application :")
    for a in [
        "le diagnostic de performance énergétique (DPE) ;",
        "l'état des risques et pollutions (ERP) ;",
        "le constat de risque d'exposition au plomb (CREP), pour les logements construits avant le 1er janvier 1949 ;",
        "l'état de l'installation intérieure d'électricité et de gaz, lorsque ces installations ont plus de 15 ans ;",
        "l'état des lieux d'entrée et l'inventaire contradictoire des meubles ;",
        "la notice d'information relative aux droits et obligations des locataires et des bailleurs.",
    ]:
        _bullet(doc, a)
    doc.add_paragraph()

    # ===== Signatures =====
    _para(doc, f"Fait à {cfg_def['lieu_signature']}, le {date_signature.strftime('%d/%m/%Y')}, en originaux dont un remis au preneur.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Le preneur\t\t\t\t\tLe bailleur").bold = True
    doc.add_paragraph()
    _para(doc, "(Faire précéder chaque signature de la mention manuscrite : « Lu et approuvé, bon pour accord »).", italic=True, size=10)

    # ===== État des lieux + Inventaire (page 2) =====
    doc.add_page_break()
    _heading(doc, "État des lieux contradictoire")
    doc.add_paragraph()
    _para(doc, "Remarques :")
    for _ in range(5):
        doc.add_paragraph()
    _para(doc, "Défauts :")
    for _ in range(5):
        doc.add_paragraph()
    _para(doc, f"Fait à {cfg_def['lieu_signature']}, le ___/___/______")
    doc.add_paragraph()

    # Inventaire détaillé
    inv_key = bien.get("inventaire_type")
    inv = bail_cfg.get("inventaires", {}).get(inv_key) if inv_key else None
    if inv:
        _heading(doc, "Inventaire contradictoire")
        doc.add_paragraph()
        labels = [("electromenager", "Électroménager"), ("vaisselle", "Vaisselle"),
                  ("linge", "Linge"), ("divers", "Divers")]
        for key, label in labels:
            items = inv.get(key) or []
            if not items:
                continue
            _para(doc, f"{label} :", bold=True)
            for item in items:
                _bullet(doc, item)
            doc.add_paragraph()
        _para(doc, f"Fait à {cfg_def['lieu_signature']}, le ___/___/______")

    return doc


def convertir_pdf(docx_path: Path) -> Path | None:
    """Tente la conversion DOCX → PDF via docx2pdf (Word) ou libreoffice. Retourne None si échec."""
    pdf_path = docx_path.with_suffix(".pdf")
    # Essai 1 : docx2pdf (Windows, Mac avec Word installé)
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception as e:
        print(f"  ! docx2pdf indisponible : {e}", file=sys.stderr)
    # Essai 2 : libreoffice headless
    try:
        import subprocess
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir",
             str(docx_path.parent), str(docx_path)],
            capture_output=True, text=True, timeout=60,
        )
        if pdf_path.exists():
            return pdf_path
        print(f"  ! libreoffice: {result.stderr[:200]}", file=sys.stderr)
    except Exception as e:
        print(f"  ! libreoffice indisponible : {e}", file=sys.stderr)
    return None


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    p = argparse.ArgumentParser(description="Génère un bail meublé (DOCX + PDF).")
    p.add_argument("--locataire", "-l", required=True, help="Nom (ou fragment) du locataire.")
    p.add_argument("--date-debut", required=True, help="Date de début du bail (YYYY-MM-DD).")
    p.add_argument("--date-signature", help="Date de signature (YYYY-MM-DD). Défaut : veille du début.")
    p.add_argument("--loyer", type=float, help="Loyer mensuel (€). Défaut : frontmatter.")
    p.add_argument("--charges", type=float, help="Charges mensuelles (€). Défaut : frontmatter.")
    p.add_argument("--depot", type=float, help="Dépôt de garantie (€). Défaut : config 1000€.")
    p.add_argument("--jour-paiement", type=int, help="Jour du mois (1-28). Défaut : config 1.")
    p.add_argument("--delai-restitution", help="Délai restitution dépôt. Défaut : '1 mois'.")
    p.add_argument("--no-pdf", action="store_true", help="Génère uniquement le DOCX, pas le PDF.")
    p.add_argument("--output-dir", help=f"Dossier sortie. Défaut : {OUTPUT_DIR}")
    p.add_argument("--dry-run", action="store_true", help="Affiche les variables sans générer.")
    args = p.parse_args()

    loc, statut = trouver_locataire(args.locataire)
    if statut == "candidat":
        print(f"  ℹ Fiche trouvée dans _Candidats/ (bail pré-signature pour {loc.nom_fichier})")
    issues = _check_fiche(loc)
    if issues and not args.dry_run:
        sous_dossier = "01. Actuels" if statut == "actuel" else "_Candidats"
        print(f"\n[BLOQUANT] Fiche locataire incomplète pour {loc.nom_fichier} :", file=sys.stderr)
        for i in issues:
            print(f"  - {i}", file=sys.stderr)
        print(f"\nComplète ces champs dans : 07. Contacts/05. Locataires/{sous_dossier}/{loc.nom_fichier}.md", file=sys.stderr)
        sys.exit(2)

    bail_cfg, biens_cfg = charger_config()
    bien = resoudre_bien(loc, biens_cfg, bail_cfg)

    if args.dry_run:
        print(f"\nLocataire : {loc.nom_fichier}")
        print(f"  Civilité : {loc.civilite}")
        print(f"  Date naissance : {loc.date_naissance}")
        print(f"  Lieu : {loc.lieu_naissance}")
        print(f"  Nationalité : {loc.nationalite}")
        print(f"\nBien : {bien['id']}")
        print(f"  Adresse : {bien['adresse_ligne1']}, {bien['complement']}, {bien['cp_ville']}")
        print(f"  Surface : {bien['surface_m2']}m²")
        print(f"  Pièces : {bien['pieces']}")
        print(f"  Charges : {bien['charges_incluses']}")
        print(f"\nBail :")
        print(f"  Date début : {args.date_debut}")
        print(f"  Loyer : {args.loyer or loc.montant_loyer}€ + {args.charges or loc.montant_charges}€ charges")
        if issues:
            print(f"\n[ATTENTION] Champs manquants : {', '.join(issues)}")
        return

    doc = render_bail(loc, bien, bail_cfg, args)

    out_dir = Path(args.output_dir) if args.output_dir else OUTPUT_DIR
    out_dir = out_dir / loc.nom_fichier
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f"Bail-{bien['id']}-{loc.nom_fichier.replace(' ', '-')}-{args.date_debut}.docx"
    docx_path = out_dir / filename
    doc.save(str(docx_path))
    print(f"  OK DOCX : {docx_path}")

    if not args.no_pdf:
        pdf_path = convertir_pdf(docx_path)
        if pdf_path:
            print(f"  OK PDF  : {pdf_path}")
        else:
            print(f"  ! PDF non généré (ni Word/docx2pdf ni LibreOffice détectés). DOCX OK.")


if __name__ == "__main__":
    main()
