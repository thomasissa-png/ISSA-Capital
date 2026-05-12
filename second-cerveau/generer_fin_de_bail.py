#!/usr/bin/env python3
"""
Génère une attestation de fin de bail (DOCX + PDF) — pour qu'un ancien locataire
puisse prouver à son assurance / sa banque / un autre bailleur qu'il a bien quitté
le logement.

Usage:
    python generer_fin_de_bail.py --locataire "Léa Lebioda" --date-fin 2024-05-17
    python generer_fin_de_bail.py -l "Elisa Morales" --date-fin 2025-08-31 --date-emission 2025-09-01
"""

from __future__ import annotations
import argparse
import datetime as dt
import os
import re
import sys
import unicodedata
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

VAULT_ROOT = Path(os.environ.get("VAULT_ROOT", r"G:\Mon Drive\Obsidian"))
LOCATAIRES_DIR_ACTUELS = VAULT_ROOT / "07. Contacts" / "05. Locataires" / "01. Actuels"
LOCATAIRES_DIR_ANCIENS = VAULT_ROOT / "07. Contacts" / "05. Locataires" / "02. Anciens"
BAUX_DIR = VAULT_ROOT / "08. Outils" / "Baux"
QUITTANCES_DIR = VAULT_ROOT / "08. Outils" / "Quittances"
BAIL_CONFIG = BAUX_DIR / "bail-config.yml"
BIENS_CONFIG = QUITTANCES_DIR / "biens.yml"
OUTPUT_DIR = Path(os.environ.get("OUTPUT_DIR_BAUX", str(BAUX_DIR / "_generes")))

MOIS_FR = ["", "janvier", "février", "mars", "avril", "mai", "juin",
           "juillet", "août", "septembre", "octobre", "novembre", "décembre"]

FRONTMATTER_RE = re.compile(r"^---\n(.*?)\n---\n", re.DOTALL)


def _parse_date(v):
    if v is None or v == "":
        return None
    if isinstance(v, dt.date):
        return v
    try:
        return dt.date.fromisoformat(str(v))
    except ValueError:
        return None


def trouver_locataire(query: str) -> tuple[str, dict]:
    """Cherche dans Actuels ET Anciens. Retourne (nom_fichier, frontmatter)."""
    def norm(s: str) -> str:
        s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode().lower()
        return re.sub(r"\s+", " ", s).strip()
    qn = norm(query)
    matches = []
    for folder in [LOCATAIRES_DIR_ACTUELS, LOCATAIRES_DIR_ANCIENS]:
        if not folder.exists():
            continue
        for f in folder.glob("*.md"):
            if f.name.startswith("_"):
                continue
            if qn in norm(f.stem):
                text = f.read_text(encoding="utf-8")
                m = FRONTMATTER_RE.match(text)
                fm = yaml.safe_load(m.group(1)) if m else {}
                matches.append((f.stem, fm or {}))
    if not matches:
        raise SystemExit(f"Locataire introuvable : {query}")
    if len(matches) > 1:
        noms = ", ".join(m[0] for m in matches)
        raise SystemExit(f"Plusieurs locataires matchent : {noms}")
    return matches[0]


def date_fr(d: dt.date) -> str:
    j = "1er" if d.day == 1 else str(d.day)
    return f"{j} {MOIS_FR[d.month]} {d.year}"


def resoudre_adresse_bien(adresse_brute: str, biens_cfg: dict) -> str:
    """Renvoie l'adresse propre du bien (sans le n° de studio) — utilisée dans le texte de l'attestation."""
    adresse_lc = adresse_brute.lower()
    for bien in biens_cfg["biens"]:
        for frag in bien["match_adresse"]:
            if frag.lower() in adresse_lc:
                return f"{bien['ligne1']}, {bien['cp_ville']}"
    # Fallback : renvoyer l'adresse brute
    return adresse_brute


def render_fin_de_bail(nom_fichier: str, fm: dict, bail_cfg: dict, biens_cfg: dict, args) -> Document:
    cfg_b = bail_cfg["bailleur"]
    cfg_def = bail_cfg["defaults"]

    nom_locataire = fm.get("nom_officiel") or nom_fichier
    adresse_bien = fm.get("adresse_bien") or ""
    adresse_propre = resoudre_adresse_bien(adresse_bien, biens_cfg)

    date_fin = dt.date.fromisoformat(args.date_fin)
    date_emission = dt.date.fromisoformat(args.date_emission) if args.date_emission else dt.date.today()
    d_naiss_bailleur = dt.date.fromisoformat(cfg_b["date_naissance"])

    doc = Document()
    sty = doc.styles["Normal"]
    sty.font.name = "Arial"
    sty.font.size = Pt(11)
    for s in doc.sections:
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)

    # En-tête bailleur (haut gauche)
    p = doc.add_paragraph()
    p.add_run(cfg_b["nom_complet"]).bold = True
    doc.add_paragraph(cfg_b["adresse"])
    doc.add_paragraph(cfg_b["cp_ville"])
    doc.add_paragraph()
    doc.add_paragraph()

    # Objet
    p = doc.add_paragraph()
    p.add_run("Objet : ").bold = True
    p.add_run("Fin de bail")
    doc.add_paragraph()

    # Lieu, date
    p = doc.add_paragraph()
    p.add_run(f"Fait à {cfg_def['lieu_signature']}, le {date_emission.strftime('%d/%m/%Y')},")
    doc.add_paragraph()

    # A qui de droit
    doc.add_paragraph("À qui de droit,")
    doc.add_paragraph()

    # Corps
    doc.add_paragraph(
        f"Je, soussigné {cfg_b['nom_complet']}, né le "
        f"{date_fr(d_naiss_bailleur)} à {cfg_b['lieu_naissance']}, "
        f"propriétaire du logement situé au {adresse_propre}, "
        f"certifie sur l'honneur que {nom_locataire} n'est plus locataire "
        f"depuis le {date_fr(date_fin)}."
    )
    doc.add_paragraph()
    doc.add_paragraph("Très cordialement,")
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature image si dispo
    sig_path_rel = cfg_b.get("signature_image", "")
    if sig_path_rel:
        sig = Path(sig_path_rel)
        if not sig.is_absolute():
            sig = BAUX_DIR / sig_path_rel
        if sig.exists():
            try:
                from docx.shared import Inches
                p = doc.add_paragraph()
                p.add_run().add_picture(str(sig), width=Inches(1.5))
            except Exception as e:
                print(f"  ! Signature non insérée : {e}", file=sys.stderr)

    doc.add_paragraph(cfg_b["nom_complet"])

    return doc


def convertir_pdf(docx_path: Path) -> Path | None:
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception:
        pass
    try:
        import subprocess
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir",
             str(docx_path.parent), str(docx_path)],
            capture_output=True, timeout=60,
        )
        if pdf_path.exists():
            return pdf_path
    except Exception:
        pass
    return None


def main():
    p = argparse.ArgumentParser(description="Génère une attestation de fin de bail.")
    p.add_argument("--locataire", "-l", required=True, help="Nom (ou fragment) du locataire.")
    p.add_argument("--date-fin", required=True, help="Date de fin du bail (YYYY-MM-DD).")
    p.add_argument("--date-emission", help="Date d'émission de l'attestation. Défaut : aujourd'hui.")
    p.add_argument("--no-pdf", action="store_true", help="Ne génère pas le PDF.")
    args = p.parse_args()

    with open(BAIL_CONFIG, "r", encoding="utf-8") as f:
        bail_cfg = yaml.safe_load(f)
    with open(BIENS_CONFIG, "r", encoding="utf-8") as f:
        biens_cfg = yaml.safe_load(f)

    nom_fichier, fm = trouver_locataire(args.locataire)
    doc = render_fin_de_bail(nom_fichier, fm, bail_cfg, biens_cfg, args)

    out_dir = OUTPUT_DIR / nom_fichier
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / f"Fin-de-bail-{nom_fichier.replace(' ', '-')}-{args.date_fin}.docx"
    doc.save(str(docx_path))
    print(f"  OK DOCX : {docx_path}")

    if not args.no_pdf:
        pdf_path = convertir_pdf(docx_path)
        if pdf_path:
            print(f"  OK PDF  : {pdf_path}")
        else:
            print(f"  ! PDF non généré. DOCX OK.")


if __name__ == "__main__":
    main()
